"""
LORENZ SaaS - RAG Service Implementation
Hybrid search with Qdrant + BM25
"""

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from typing import Optional, List, Tuple
from uuid import UUID
import hashlib
import logging

from app.models import User, RAGDocument
from app.config import settings

logger = logging.getLogger(__name__)


class RAGService:
    """RAG (Retrieval Augmented Generation) service"""

    def __init__(self, db: AsyncSession):
        self.db = db
        self._qdrant_client = None
        self._embedder = None

    async def _get_qdrant_client(self):
        """Lazy load Qdrant client"""
        if self._qdrant_client is None:
            try:
                from qdrant_client import QdrantClient
                self._qdrant_client = QdrantClient(
                    host=settings.QDRANT_HOST,
                    port=settings.QDRANT_PORT,
                    api_key=settings.QDRANT_API_KEY
                )
            except Exception as e:
                logger.warning(f"Could not connect to Qdrant: {e}")
                return None
        return self._qdrant_client

    async def _get_embedder(self):
        """Lazy load sentence transformer"""
        if self._embedder is None:
            try:
                from sentence_transformers import SentenceTransformer
                self._embedder = SentenceTransformer('all-MiniLM-L6-v2')
            except Exception as e:
                logger.warning(f"Could not load embedder: {e}")
                return None
        return self._embedder

    async def search(
        self,
        user_id: UUID,
        tenant_id: UUID,
        query: str,
        source_types: Optional[List[str]] = None,
        limit: int = 10
    ) -> List[dict]:
        """
        Hybrid search across knowledge base.
        Combines semantic search (Qdrant) with keyword search (BM25).
        """
        results = []

        # 1. Semantic search via Qdrant
        embedder = await self._get_embedder()
        qdrant = await self._get_qdrant_client()

        if embedder and qdrant:
            try:
                query_embedding = embedder.encode(query).tolist()
                collection_name = f"tenant_{tenant_id}"

                qdrant_results = qdrant.search(
                    collection_name=collection_name,
                    query_vector=query_embedding,
                    limit=limit * 2,
                    query_filter={
                        "must": [
                            {"key": "user_id", "match": {"value": str(user_id)}}
                        ]
                    } if source_types is None else {
                        "must": [
                            {"key": "user_id", "match": {"value": str(user_id)}},
                            {"key": "source_type", "match": {"any": source_types}}
                        ]
                    }
                )

                for hit in qdrant_results:
                    results.append({
                        "id": hit.id,
                        "score": hit.score,
                        "source": "semantic",
                        **hit.payload
                    })
            except Exception as e:
                logger.warning(f"Qdrant search failed: {e}")

        # 2. Keyword search from database (simple LIKE for now)
        # In production, use PostgreSQL full-text search or dedicated BM25
        db_query = select(RAGDocument).where(
            RAGDocument.user_id == user_id,
            RAGDocument.content.ilike(f"%{query}%")
        )

        if source_types:
            db_query = db_query.where(RAGDocument.source_type.in_(source_types))

        db_query = db_query.limit(limit)
        result = await self.db.execute(db_query)
        db_docs = result.scalars().all()

        for doc in db_docs:
            # Check if already in results
            if not any(r.get("id") == str(doc.id) for r in results):
                results.append({
                    "id": str(doc.id),
                    "score": 0.5,  # Base score for keyword matches
                    "source": "keyword",
                    "title": doc.title,
                    "snippet": doc.content[:300] if doc.content else "",
                    "source_type": doc.source_type,
                    "metadata": doc.metadata
                })

        # 3. Reciprocal Rank Fusion to combine results
        results = self._rrf_merge(results)[:limit]

        return results

    def _rrf_merge(self, results: List[dict], k: int = 60) -> List[dict]:
        """Merge results using Reciprocal Rank Fusion"""
        scores = {}
        docs = {}

        for rank, result in enumerate(results):
            doc_id = result["id"]
            rrf_score = 1 / (k + rank + 1)

            if doc_id in scores:
                scores[doc_id] += rrf_score
            else:
                scores[doc_id] = rrf_score
                docs[doc_id] = result

        # Sort by combined score
        sorted_ids = sorted(scores.keys(), key=lambda x: scores[x], reverse=True)

        return [
            {**docs[doc_id], "combined_score": scores[doc_id]}
            for doc_id in sorted_ids
        ]

    async def index_document(
        self,
        user_id: UUID,
        tenant_id: UUID,
        filename: str,
        content: bytes,
        content_type: str
    ):
        """Index an uploaded document"""
        # Extract text based on content type
        text = await self._extract_text(content, content_type)

        if not text:
            logger.warning(f"Could not extract text from {filename}")
            return

        # Compute content hash for deduplication
        content_hash = hashlib.sha256(text.encode()).hexdigest()

        # Check for duplicate
        query = select(RAGDocument).where(
            RAGDocument.user_id == user_id,
            RAGDocument.content_hash == content_hash
        )
        result = await self.db.execute(query)
        if result.scalar_one_or_none():
            logger.info(f"Document already indexed: {filename}")
            return

        # Chunk the text
        chunks = self._chunk_text(text)

        # Create embeddings and store
        embedder = await self._get_embedder()
        qdrant = await self._get_qdrant_client()

        for i, chunk in enumerate(chunks):
            # Create database record
            doc = RAGDocument(
                user_id=user_id,
                source_type="file",
                title=filename,
                content=chunk,
                content_hash=content_hash if i == 0 else None,
                chunk_index=i,
                total_chunks=len(chunks),
                metadata={
                    "content_type": content_type,
                    "filename": filename
                },
                status="indexed"
            )
            self.db.add(doc)
            await self.db.flush()

            # Store embedding in Qdrant
            if embedder and qdrant:
                try:
                    embedding = embedder.encode(chunk).tolist()
                    collection_name = f"tenant_{tenant_id}"

                    # Ensure collection exists
                    await self._ensure_collection(qdrant, collection_name)

                    qdrant.upsert(
                        collection_name=collection_name,
                        points=[{
                            "id": str(doc.id),
                            "vector": embedding,
                            "payload": {
                                "user_id": str(user_id),
                                "source_type": "file",
                                "title": filename,
                                "snippet": chunk[:300],
                                "metadata": doc.metadata
                            }
                        }]
                    )

                    doc.qdrant_point_id = str(doc.id)
                    self.db.add(doc)
                except Exception as e:
                    logger.error(f"Failed to store embedding: {e}")

        await self.db.commit()
        logger.info(f"Indexed document: {filename} ({len(chunks)} chunks)")

    async def _ensure_collection(self, qdrant, collection_name: str):
        """Ensure Qdrant collection exists"""
        try:
            qdrant.get_collection(collection_name)
        except Exception:
            from qdrant_client.models import Distance, VectorParams
            qdrant.create_collection(
                collection_name=collection_name,
                vectors_config=VectorParams(size=384, distance=Distance.COSINE)
            )

    async def _extract_text(self, content: bytes, content_type: str) -> Optional[str]:
        """Extract text from document"""
        if content_type == "text/plain" or content_type == "text/markdown":
            return content.decode("utf-8", errors="ignore")

        elif content_type == "application/pdf":
            try:
                import PyPDF2
                import io
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
                return None

        elif "wordprocessingml" in content_type:  # DOCX
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(content))
                text = "\n".join([p.text for p in doc.paragraphs])
                return text
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                return None

        return None

    def _chunk_text(self, text: str, max_tokens: int = 500) -> List[str]:
        """Split text into chunks"""
        # Simple sentence-based chunking
        # In production, use more sophisticated chunking
        sentences = text.replace("\n", " ").split(". ")
        chunks = []
        current_chunk = ""

        for sentence in sentences:
            if len(current_chunk) + len(sentence) > max_tokens * 4:  # ~4 chars per token
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                current_chunk += ". " + sentence if current_chunk else sentence

        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks if chunks else [text[:2000]]

    async def delete_document(
        self,
        document_id: UUID,
        user_id: UUID,
        tenant_id: UUID
    ):
        """Delete a document from knowledge base"""
        query = select(RAGDocument).where(
            RAGDocument.id == document_id,
            RAGDocument.user_id == user_id
        )
        result = await self.db.execute(query)
        doc = result.scalar_one_or_none()

        if not doc:
            raise ValueError("Document not found")

        # Remove from Qdrant
        qdrant = await self._get_qdrant_client()
        if qdrant and doc.qdrant_point_id:
            try:
                qdrant.delete(
                    collection_name=f"tenant_{tenant_id}",
                    points_selector=[doc.qdrant_point_id]
                )
            except Exception as e:
                logger.warning(f"Failed to delete from Qdrant: {e}")

        await self.db.delete(doc)
        await self.db.commit()

    async def list_documents(
        self,
        user_id: UUID,
        source_type: Optional[str] = None,
        limit: int = 50,
        offset: int = 0
    ) -> Tuple[List[dict], int]:
        """List indexed documents"""
        query = select(RAGDocument).where(
            RAGDocument.user_id == user_id,
            RAGDocument.chunk_index == 0  # Only parent documents
        )

        if source_type:
            query = query.where(RAGDocument.source_type == source_type)

        # Count
        count_query = select(func.count()).select_from(query.subquery())
        count_result = await self.db.execute(count_query)
        total = count_result.scalar()

        # Get paginated
        query = query.order_by(RAGDocument.created_at.desc()).offset(offset).limit(limit)
        result = await self.db.execute(query)
        documents = result.scalars().all()

        return [
            {
                "id": str(d.id),
                "title": d.title,
                "source_type": d.source_type,
                "total_chunks": d.total_chunks,
                "status": d.status,
                "metadata": d.metadata,
                "created_at": d.created_at
            }
            for d in documents
        ], total

    async def get_source_summary(self, user_id: UUID) -> List[dict]:
        """Get summary of indexed sources"""
        query = select(
            RAGDocument.source_type,
            func.count(RAGDocument.id)
        ).where(
            RAGDocument.user_id == user_id
        ).group_by(RAGDocument.source_type)

        result = await self.db.execute(query)
        rows = result.all()

        return [
            {"source_type": row[0], "count": row[1]}
            for row in rows
        ]

    async def index_email(
        self,
        user_id: UUID,
        tenant_id: UUID,
        email_data: dict
    ) -> Optional[str]:
        """
        Index an email message into the RAG knowledge base.

        Args:
            user_id: User ID
            tenant_id: Tenant ID
            email_data: Dict containing email fields:
                - message_id: Unique message ID
                - from_address: Sender email
                - to_addresses: List of recipients
                - subject: Email subject
                - body: Email body text
                - date: Email date
                - attachments: Optional list of attachment names

        Returns:
            Document ID if indexed, None if skipped (duplicate)
        """
        from datetime import datetime

        message_id = email_data.get("message_id", "")
        subject = email_data.get("subject", "")
        body = email_data.get("body", "")
        from_address = email_data.get("from_address", "")
        to_addresses = email_data.get("to_addresses", [])
        date = email_data.get("date")
        attachments = email_data.get("attachments", [])

        if not body or len(body.strip()) < 10:
            logger.debug(f"Skipping email with minimal content: {subject}")
            return None

        # Create content hash for deduplication
        content_hash = hashlib.sha256(
            f"{message_id}:{subject}:{body[:500]}".encode()
        ).hexdigest()

        # Check for duplicate
        query = select(RAGDocument).where(
            RAGDocument.user_id == user_id,
            RAGDocument.source_type == "email",
            RAGDocument.content_hash == content_hash
        )
        result = await self.db.execute(query)
        if result.scalar_one_or_none():
            logger.debug(f"Email already indexed: {subject}")
            return None

        # Prepare email content for indexing
        email_text = f"""
Subject: {subject}
From: {from_address}
To: {', '.join(to_addresses) if isinstance(to_addresses, list) else to_addresses}
Date: {date}

{body}
"""

        # Chunk the email if it's long
        chunks = self._chunk_text(email_text, max_tokens=400)

        # Get embedder and Qdrant
        embedder = await self._get_embedder()
        qdrant = await self._get_qdrant_client()

        first_doc_id = None

        for i, chunk in enumerate(chunks):
            # Create database record
            doc = RAGDocument(
                user_id=user_id,
                source_type="email",
                title=f"{subject[:100]}",
                content=chunk,
                content_hash=content_hash if i == 0 else None,
                chunk_index=i,
                total_chunks=len(chunks),
                metadata={
                    "message_id": message_id,
                    "from_address": from_address,
                    "to_addresses": to_addresses,
                    "subject": subject,
                    "date": str(date) if date else None,
                    "has_attachments": len(attachments) > 0,
                    "attachment_names": attachments[:5],  # Limit to 5
                },
                status="indexed"
            )
            self.db.add(doc)
            await self.db.flush()

            if i == 0:
                first_doc_id = str(doc.id)

            # Store embedding in Qdrant
            if embedder and qdrant:
                try:
                    embedding = embedder.encode(chunk).tolist()
                    collection_name = f"tenant_{tenant_id}"

                    # Ensure collection exists
                    await self._ensure_collection(qdrant, collection_name)

                    qdrant.upsert(
                        collection_name=collection_name,
                        points=[{
                            "id": str(doc.id),
                            "vector": embedding,
                            "payload": {
                                "user_id": str(user_id),
                                "source_type": "email",
                                "title": subject[:100],
                                "snippet": chunk[:300],
                                "from_address": from_address,
                                "date": str(date) if date else None,
                                "metadata": doc.metadata
                            }
                        }]
                    )

                    doc.qdrant_point_id = str(doc.id)
                    self.db.add(doc)
                except Exception as e:
                    logger.error(f"Failed to store email embedding: {e}")

        await self.db.commit()
        logger.info(f"Indexed email: {subject[:50]} ({len(chunks)} chunks)")

        return first_doc_id

    async def index_emails_batch(
        self,
        user_id: UUID,
        tenant_id: UUID,
        emails: List[dict]
    ) -> dict:
        """
        Batch index multiple emails.

        Returns:
            dict with indexed count, skipped count, and errors
        """
        indexed = 0
        skipped = 0
        errors = []

        for email_data in emails:
            try:
                result = await self.index_email(user_id, tenant_id, email_data)
                if result:
                    indexed += 1
                else:
                    skipped += 1
            except Exception as e:
                logger.error(f"Failed to index email: {e}")
                errors.append({
                    "subject": email_data.get("subject", "Unknown"),
                    "error": str(e)
                })

        logger.info(f"Batch indexed {indexed} emails, skipped {skipped}, errors {len(errors)}")

        return {
            "indexed": indexed,
            "skipped": skipped,
            "errors": errors
        }

    async def reindex_user_content(
        self,
        user_id: UUID,
        tenant_id: UUID,
        source_type: Optional[str] = None
    ):
        """Reindex all user content"""
        # Placeholder - implement full reindex logic
        logger.info(f"Reindexing content for user {user_id}")

    async def get_stats(self, user_id: UUID) -> dict:
        """Get RAG statistics"""
        query = select(func.count(RAGDocument.id)).where(
            RAGDocument.user_id == user_id
        )
        result = await self.db.execute(query)
        total_docs = result.scalar()

        sources = await self.get_source_summary(user_id)

        return {
            "total_documents": total_docs,
            "sources": sources
        }
